# 标准库
import os
import re
import shutil
from datetime import datetime
from typing import List, Dict, Any

# 第三方包
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PROGRAM_DIR = os.path.join(BASE_DIR, "market_program")
ARCHIVE_DIR = os.path.join(PROGRAM_DIR, ".archive")
MAX_ARCHIVE = 3


def _ensure_dirs():
    os.makedirs(PROGRAM_DIR, exist_ok=True)
    os.makedirs(ARCHIVE_DIR, exist_ok=True)


def _set_font(run, name='微软雅黑', size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_font(run, size=16 if level == 1 else 14, bold=True)


def _add_content(doc, text):
    if not text:
        return
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        run = p.add_run(line)
        _set_font(run, size=11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5


def _add_cover(doc, project_name):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('营销战略方案')
    _set_font(run, size=28, bold=True, color=RGBColor(0x1A, 0x56, 0xDB))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(project_name)
    _set_font(run2, size=18, color=RGBColor(0x55, 0x55, 0x55))

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
    _set_font(run3, size=10, color=RGBColor(0x88, 0x88, 0x88))

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('虚拟营销团队 · 多Agent协作系统')
    _set_font(run4, size=10, color=RGBColor(0x88, 0x88, 0x88))

    doc.add_page_break()


def _strip_markdown(text: str) -> str:
    """移除 LLM 输出中常见的 markdown 格式符号。"""
    # 加粗 **text** -> text
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # 斜体 *text*（但不匹配列表项）
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', text)
    # 行内代码 `text` -> text
    text = re.sub(r'`(.+?)`', r'\1', text)
    # 标题标记 # ## ### 等
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    # 无序列表标记 - * +
    text = re.sub(r'^[\-\*\+]\s+', '', text, flags=re.MULTILINE)
    # 有序列表标记 1. 2. 等
    text = re.sub(r'^\d+\.\s+', '', text, flags=re.MULTILINE)
    # 水平线 --- *** ===
    text = re.sub(r'^[\-\*\_]{3,}\s*$', '', text, flags=re.MULTILINE)
    # 引用 >
    text = re.sub(r'^>\s+', '', text, flags=re.MULTILINE)
    return text.strip()


def _format_kv(doc, key: str, val):
    """递归格式化单个 key-value 对到 docx。"""
    label = _field_label(key)
    if isinstance(val, list):
        if val and isinstance(val[0], dict):
            # list of dicts —— 如 strategies / ideas / copies
            p = doc.add_paragraph()
            run = p.add_run(label)
            _set_font(run, size=12, bold=True)
            for i, item in enumerate(val):
                title_key = 'campaign_name' if 'campaign_name' in item else 'name'
                item_title = item.get(title_key, f'#{i+1}')
                p2 = doc.add_paragraph()
                run2 = p2.add_run(f'  {item_title}')
                _set_font(run2, size=11, bold=True)
                for ik, iv in item.items():
                    if ik in ('name', 'campaign_name'):
                        continue
                    _format_kv(doc, ik, iv)
        else:
            # list of strings
            p = doc.add_paragraph()
            run = p.add_run(label)
            _set_font(run, size=11, bold=True)
            for item in val:
                bp = doc.add_paragraph(_strip_markdown(str(item)), style='List Bullet')
                for r in bp.runs:
                    _set_font(r, size=11)
    elif isinstance(val, bool):
        p = doc.add_paragraph()
        status = "✅ 通过" if val else "❌ 未通过"
        run = p.add_run(f"{label}：{status}")
        _set_font(run, size=11)
    elif isinstance(val, dict):
        p = doc.add_paragraph()
        run = p.add_run(label)
        _set_font(run, size=12, bold=True)
        for k, v in val.items():
            _format_kv(doc, k, v)
    else:
        text = _strip_markdown(str(val))
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：{text}")
        _set_font(run, size=11)
        p.paragraph_format.space_after = Pt(4)


def _format_output(doc, output_dict: dict):
    """将 Pydantic 输出 dict 格式化为可读文本，对拍平策略字段自动分组。"""
    if not output_dict:
        return

    # 检测拍平的 MarketStrategy 字段并分组渲染
    if 'strategy_a_name' in output_dict:
        p = doc.add_paragraph()
        run = p.add_run('备选方案')
        _set_font(run, size=12, bold=True)
        for group_key in ('a', 'b', 'c'):
            name = output_dict.get(f'strategy_{group_key}_name', '')
            tactics = output_dict.get(f'strategy_{group_key}_tactics', [])
            channels = output_dict.get(f'strategy_{group_key}_channels', [])
            kpis = output_dict.get(f'strategy_{group_key}_kpis', [])
            # 方案标题
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f'  {name}')
            _set_font(run2, size=11, bold=True)
            # 战术
            if tactics:
                p3 = doc.add_paragraph()
                run3 = p3.add_run('    战术清单：')
                _set_font(run3, size=11, bold=True)
                for t in tactics:
                    bp = doc.add_paragraph(_strip_markdown(str(t)), style='List Bullet')
                    for r in bp.runs:
                        _set_font(r, size=11)
            # 渠道
            if channels:
                p4 = doc.add_paragraph()
                run4 = p4.add_run('    渠道清单：')
                _set_font(run4, size=11, bold=True)
                for c in channels:
                    bp = doc.add_paragraph(_strip_markdown(str(c)), style='List Bullet')
                    for r in bp.runs:
                        _set_font(r, size=11)
            # KPI
            if kpis:
                p5 = doc.add_paragraph()
                run5 = p5.add_run('    关键绩效指标(KPI)：')
                _set_font(run5, size=11, bold=True)
                for k in kpis:
                    bp = doc.add_paragraph(_strip_markdown(str(k)), style='List Bullet')
                    for r in bp.runs:
                        _set_font(r, size=11)

        rec = output_dict.get('recommendation', '')
        if rec:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(f'推荐方案：{_strip_markdown(str(rec))}')
            _set_font(run, size=11)
        return

    # 通用渲染（跳过已回流到活动创意/文案的修正版本）
    for key, val in output_dict.items():
        if key in ('corrected_ideas', 'corrected_copies'):
            continue
        _format_kv(doc, key, val)


def _field_label(key: str) -> str:
    labels = {
        'requirement_analysis': '需求分析',
        'task_breakdown': '任务拆解',
        'key_guidance': '关键指引',
        'market_overview': '市场概况',
        'competitor_analysis': '竞品分析',
        'audience_insights': '受众洞察',
        'strategy_a_name': '方案A', 'strategy_a_tactics': '战术清单', 'strategy_a_channels': '渠道清单', 'strategy_a_kpis': 'KPI',
        'strategy_b_name': '方案B', 'strategy_b_tactics': '战术清单', 'strategy_b_channels': '渠道清单', 'strategy_b_kpis': 'KPI',
        'strategy_c_name': '方案C', 'strategy_c_tactics': '战术清单', 'strategy_c_channels': '渠道清单', 'strategy_c_kpis': 'KPI',
        'recommendation': '推荐方案',
        'ideas': '活动创意',
        'description': '描述',
        'audience': '目标受众',
        'channel': '渠道',
        'copies': '营销文案',
        'campaign_name': '对应活动',
        'title': '标题',
        'body': '正文',
        'approved': '审核结论',
        'issues_found': '发现的问题',
        'revisions_made': '修正对照',
        'corrected_ideas': '修正后活动创意',
        'corrected_copies': '修正后文案',
        'review_notes': '审核备注',
    }
    return labels.get(key, key)


def generate_program_docx(project_name: str, agent_outputs: List[Dict[str, Any]]) -> str:
    """根据 Agent 输出生成格式化方案文档，返回文件路径。"""
    _ensure_dirs()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)

    _add_cover(doc, project_name)

    # 按 agent_key 映射到章节
    section_map = {
        'pm_output': ('需求分析', '项目经理对用户需求的分析和任务拆解'),
        'research_report': ('市场调研报告', '市场分析师完成的市场、竞品和受众调研'),
        'market_strategy': ('营销战略方案', '营销战略师制定的战略框架（目标、战术、渠道、KPI）'),
        'campaign_ideas': ('活动创意', '内容创作师策划的营销活动创意'),
        'copy': ('营销文案', '内容创作师撰写的具体营销文案'),
        'review_result': ('审核意见', '内容审核员的质量审核结论与修正说明'),
    }

    for item in agent_outputs:
        key = item.get('agent_key', '')
        output = item.get('output', {})

        if key not in section_map:
            continue

        section_title, section_desc = section_map[key]
        _add_section_heading(doc, section_title, level=1)

        p = doc.add_paragraph()
        run = p.add_run(section_desc)
        _set_font(run, size=9, color=RGBColor(0x99, 0x99, 0x99))
        p.paragraph_format.space_after = Pt(8)

        if isinstance(output, dict):
            _format_output(doc, output)
        else:
            _add_content(doc, str(output))

        doc.add_paragraph()  # 章节间距

    # 总结
    _add_section_heading(doc, '总结与下一步行动', level=1)
    _add_content(doc, f'本方案由虚拟营销团队自动生成，包含从需求分析到内容审核的完整营销战略设计。'
                      f'建议下一步由人工进行最终审核和优化调整。')

    # 保存文件
    safe_name = re.sub(r'[\\/:*?"<>|]', '', project_name)
    filename = f"{safe_name}.docx"
    filepath = os.path.join(PROGRAM_DIR, filename)

    # 如果已有同名文件，备份到 .archive/
    if os.path.exists(filepath):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_name = f"{safe_name}_{timestamp}.docx"
        shutil.copy2(filepath, os.path.join(ARCHIVE_DIR, backup_name))
        # 只保留最近 3 个备份
        backups = sorted(
            [f for f in os.listdir(ARCHIVE_DIR) if f.startswith(safe_name)],
            reverse=True
        )
        for old in backups[MAX_ARCHIVE:]:
            os.remove(os.path.join(ARCHIVE_DIR, old))

    doc.save(filepath)
    return filepath
